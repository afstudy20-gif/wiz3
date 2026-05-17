"""Publication export: Word document tables and figure captions."""
import io
import json
import numpy as np
import pandas as pd
from typing import Dict, List, Optional

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from services import store

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

router = APIRouter()


def _get_df(session_id: str) -> pd.DataFrame:
    df = store.get_filtered(session_id)
    if df is None:
        raise HTTPException(status_code=404, detail="Session not found")
    return df


# ── Table 1 DOCX export ──────────────────────────────────────────────────────

class TableDocxRequest(BaseModel):
    session_id: str
    group_column: Optional[str] = None
    variables: List[str]
    variable_kinds: Optional[Dict[str, str]] = None
    selected_stats: Optional[List[str]] = None


def _set_cell_border(cell, **kwargs):
    """Set cell border properties (top, bottom, left, right)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = tcPr.makeelement(qn("w:tcBorders"), {})
        tcPr.append(tcBorders)
    for edge, val in kwargs.items():
        element = tcBorders.find(qn(f"w:{edge}"))
        if element is None:
            element = tcBorders.makeelement(qn(f"w:{edge}"), {})
            tcBorders.append(element)
        element.set(qn("w:val"), val.get("val", "single"))
        element.set(qn("w:sz"), val.get("sz", "4"))
        element.set(qn("w:color"), val.get("color", "000000"))
        element.set(qn("w:space"), val.get("space", "0"))


def _run_table1_analysis(req: TableDocxRequest) -> dict:
    """Run the Table 1 analysis (mirrors stats.table1 logic)."""
    from scipy import stats as scipy_stats

    df = _get_df(req.session_id)
    rows = []
    sel_stats = req.selected_stats if req.selected_stats else ["auto"]

    groups = None
    group_labels = []
    group_ns: dict = {}
    if req.group_column and req.group_column in df.columns:
        groups = sorted(df[req.group_column].dropna().unique().tolist(), key=str)
        group_labels = [str(g) for g in groups]
        group_ns = {str(g): int((df[req.group_column] == g).sum()) for g in groups}

    def _f(v, d=2):
        if v is None or (isinstance(v, float) and (np.isnan(v) or np.isinf(v))):
            return "\u2014"
        return f"{v:.{d}f}"

    for var in req.variables:
        if var not in df.columns:
            continue
        s = df[var]
        provided_kind = (req.variable_kinds or {}).get(var)
        if provided_kind == "numeric":
            is_num = True
        elif provided_kind in ("categorical", "text", "boolean"):
            is_num = False
        else:
            is_num = pd.api.types.is_numeric_dtype(s) and s.nunique() > 10

        if is_num:
            s_all = s.dropna().astype(float)
            n = len(s_all)
            normal = True
            if 3 <= n <= 2000:
                _, p_norm = scipy_stats.shapiro(s_all[:5000])
                normal = p_norm > 0.05
            elif n > 2000:
                skewness = float(scipy_stats.skew(s_all))
                normal = abs(skewness) <= 1.5

            if normal:
                overall = f"{_f(s_all.mean())} \u00b1 {_f(s_all.std())}"
            else:
                q1, q3 = s_all.quantile(0.25), s_all.quantile(0.75)
                overall = f"{_f(s_all.median())} [{_f(q1)}\u2013{_f(q3)}]"

            grp_vals = {}
            test_name = None
            p_str = None
            if groups is not None:
                arrs = []
                for g, gl in zip(groups, group_labels):
                    gs = df[df[req.group_column] == g][var].dropna().astype(float)
                    arrs.append(gs)
                    if normal:
                        grp_vals[gl] = f"{_f(gs.mean())} \u00b1 {_f(gs.std())}"
                    else:
                        gq1, gq3 = gs.quantile(0.25), gs.quantile(0.75)
                        grp_vals[gl] = f"{_f(gs.median())} [{_f(gq1)}\u2013{_f(gq3)}]"
                if len(arrs) >= 2:
                    try:
                        if len(groups) == 2:
                            if normal:
                                _, p_t = scipy_stats.ttest_ind(*arrs, equal_var=False)
                                test_name = "t-test"
                            else:
                                _, p_t = scipy_stats.mannwhitneyu(*arrs, alternative="two-sided")
                                test_name = "Mann\u2013Whitney"
                        else:
                            if normal:
                                _, p_t = scipy_stats.f_oneway(*arrs)
                                test_name = "ANOVA"
                            else:
                                _, p_t = scipy_stats.kruskal(*arrs)
                                test_name = "Kruskal\u2013Wallis"
                        p_str = "<0.001" if p_t < 0.001 else f"{p_t:.3f}"
                    except Exception:
                        p_str = "N/A"

            stat_label = "Mean \u00b1 SD" if normal else "Median [IQR]"
            rows.append({
                "variable": var, "type": "numeric",
                "stat_label": stat_label, "overall": overall,
                "group_stats": grp_vals, "p_value": p_str,
                "test": test_name, "sub_rows": [],
            })
        else:
            vc = s.value_counts(dropna=True)
            total = s.count()
            cats = [str(v) for v in vc.index.tolist()]
            sub_rows = []
            for cat in cats:
                n_all = int((s.astype(str) == cat).sum())
                pct = round(n_all / total * 100, 1) if total else 0
                sub = {"category": cat, "overall": f"{n_all} ({pct}%)"}
                sub["group_stats"] = {}
                if groups is not None:
                    for g, gl in zip(groups, group_labels):
                        gs = df[df[req.group_column] == g][var]
                        ng = int((gs.astype(str) == cat).sum())
                        tg = gs.count()
                        pg = round(ng / tg * 100, 1) if tg else 0
                        sub["group_stats"][gl] = f"{ng} ({pg}%)"
                sub_rows.append(sub)

            p_str = None
            test_name = None
            if groups is not None:
                try:
                    ct = pd.crosstab(df[var].astype(str), df[req.group_column])
                    chi2, p_chi, dof, expected = scipy_stats.chi2_contingency(ct)
                    if ct.shape == (2, 2) and (expected < 5).any():
                        _, p_chi = scipy_stats.fisher_exact(ct.values)
                        test_name = "Fisher"
                    else:
                        test_name = "Chi-square"
                    p_str = "<0.001" if p_chi < 0.001 else f"{p_chi:.3f}"
                except Exception:
                    p_str = "N/A"

            rows.append({
                "variable": var, "type": "categorical",
                "stat_label": "n (%)", "overall": f"n={total}",
                "group_stats": {}, "p_value": p_str,
                "test": test_name, "sub_rows": sub_rows,
            })

    return {
        "group_column": req.group_column,
        "group_labels": group_labels,
        "group_ns": group_ns,
        "total_n": len(df),
        "rows": rows,
    }


@router.post("/table_docx")
async def table_docx(req: TableDocxRequest):
    """Generate Table 1 as a Word document (.docx) download."""
    if not HAS_DOCX:
        raise HTTPException(
            status_code=501,
            detail="python-docx is not installed. Run: pip install python-docx",
        )

    result = _run_table1_analysis(req)
    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Table 1. Baseline Characteristics")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # Determine columns
    has_groups = bool(result["group_labels"])
    col_headers = ["Characteristic"]
    col_headers.append(f"Overall (N={result['total_n']})")
    if has_groups:
        for gl in result["group_labels"]:
            n = result["group_ns"].get(gl, "")
            col_headers.append(f"{gl} (n={n})")
    col_headers.append("P value")

    n_cols = len(col_headers)

    # Create table
    table = doc.add_table(rows=1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, text in enumerate(col_headers):
        cell = hdr.cells[i]
        cell.text = text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "Times New Roman"

    # Collect footnotes for test methods
    footnotes = {}

    for row_data in result["rows"]:
        if row_data["type"] == "numeric":
            cells = [f"{row_data['variable']}, {row_data['stat_label']}"]
            cells.append(row_data["overall"])
            if has_groups:
                for gl in result["group_labels"]:
                    cells.append(row_data["group_stats"].get(gl, "\u2014"))
            p_val = row_data.get("p_value", "")
            cells.append(p_val or "")
            if row_data.get("test"):
                footnotes[row_data["test"]] = True

            row_cells = table.add_row().cells
            for i, text in enumerate(cells):
                row_cells[i].text = str(text) if text else ""
                for paragraph in row_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = "Times New Roman"

        else:
            # Category header row
            cells = [f"{row_data['variable']}, n (%)"]
            cells.append(row_data["overall"])
            if has_groups:
                cells.extend([""] * len(result["group_labels"]))
            p_val = row_data.get("p_value", "")
            cells.append(p_val or "")
            if row_data.get("test"):
                footnotes[row_data["test"]] = True

            row_cells = table.add_row().cells
            for i, text in enumerate(cells):
                row_cells[i].text = str(text) if text else ""
                for paragraph in row_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = "Times New Roman"

            # Sub-rows for each category
            for sub in row_data.get("sub_rows", []):
                cells = [f"  {sub['category']}"]
                cells.append(sub["overall"])
                if has_groups:
                    for gl in result["group_labels"]:
                        cells.append(sub["group_stats"].get(gl, "\u2014"))
                cells.append("")

                row_cells = table.add_row().cells
                for i, text in enumerate(cells):
                    row_cells[i].text = str(text) if text else ""
                    for paragraph in row_cells[i].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                            run.font.name = "Times New Roman"

    # Footnotes
    if footnotes:
        fn_text = "Statistical tests: " + ", ".join(sorted(footnotes.keys())) + "."
        fn_para = doc.add_paragraph()
        fn_run = fn_para.add_run(fn_text)
        fn_run.font.size = Pt(9)
        fn_run.font.name = "Times New Roman"
        fn_run.italic = True

    # Save to buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        iter([buf.getvalue()]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="Table1.docx"'},
    )


# ── Figure Caption Generator ────────────────────────────────────────────────

class FigureCaptionRequest(BaseModel):
    figure_type: str  # "roc", "km", "forest", "histogram", "scatter", "bar"
    params: Dict = {}


_FIGURE_TEMPLATES = {
    "roc": {
        "caption": "Figure {n}. Receiver operating characteristic curve for {outcome}.",
        "footnote": "AUC = {auc} (95% CI: {ci_lower}\u2013{ci_upper}). Optimal cutoff at {cutoff} (sensitivity = {sensitivity}, specificity = {specificity}).",
    },
    "km": {
        "caption": "Figure {n}. Kaplan\u2013Meier survival curves for {outcome} stratified by {group}.",
        "footnote": "Log-rank test P = {p_value}. Median survival: {median_survival}.",
    },
    "forest": {
        "caption": "Figure {n}. Forest plot of {analysis_type} for {outcome}.",
        "footnote": "Effect sizes shown with 95% confidence intervals. I\u00b2 = {i_squared}%.",
    },
    "histogram": {
        "caption": "Figure {n}. Distribution of {variable}.",
        "footnote": "N = {n_obs}. Shapiro\u2013Wilk P = {normality_p}.",
    },
    "scatter": {
        "caption": "Figure {n}. Scatter plot of {x_var} vs {y_var}.",
        "footnote": "Pearson r = {r_value}, P = {p_value}. N = {n_obs}.",
    },
    "bar": {
        "caption": "Figure {n}. Distribution of {variable} across {group}.",
        "footnote": "Values shown as counts with percentages.",
    },
}


@router.post("/figure_caption")
async def figure_caption(req: FigureCaptionRequest):
    """Generate a publication-ready figure caption and footnote."""
    fig_type = req.figure_type.lower()
    template = _FIGURE_TEMPLATES.get(fig_type)

    if template is None:
        # Generic fallback
        caption = f"Figure 1. {fig_type.replace('_', ' ').title()} plot."
        footnote = ""
    else:
        params = {"n": 1, **req.params}
        try:
            caption = template["caption"].format_map(
                {k: params.get(k, f"[{k}]") for k in _extract_keys(template["caption"])}
            )
        except Exception:
            caption = template["caption"]
        try:
            footnote = template["footnote"].format_map(
                {k: params.get(k, f"[{k}]") for k in _extract_keys(template["footnote"])}
            )
        except Exception:
            footnote = template["footnote"]

    return {"caption": caption, "footnote": footnote}


def _extract_keys(template: str) -> list:
    """Extract {key} placeholders from a format string."""
    import re
    return re.findall(r"\{(\w+)\}", template)


# ── Method appendix DOCX from session audit log ──────────────────────────────

# Map raw audit action codes to human-readable Methods-section phrasing.
_ACTION_HUMAN: Dict[str, str] = {
    "ttest": "Independent samples t-test (Welch / pooled)",
    "anova": "One-way ANOVA with Tukey HSD post-hoc",
    "ancova": "Analysis of covariance (ANCOVA)",
    "two_way_anova": "Two-way ANOVA",
    "rm_anova": "Repeated-measures ANOVA",
    "mixed_anova": "Mixed-design ANOVA",
    "paired_ttest": "Paired-samples t-test",
    "mannwhitney": "Mann-Whitney U test",
    "wilcoxon_signed_rank": "Wilcoxon signed-rank test",
    "kruskal": "Kruskal-Wallis H test with Dunn post-hoc",
    "friedman": "Friedman test with pairwise Wilcoxon (Holm-corrected)",
    "chisquare": "Chi-square test of independence",
    "fisher": "Fisher's exact test",
    "mcnemar": "McNemar's test",
    "cochran_q": "Cochran's Q test",
    "mantel_haenszel": "Cochran-Mantel-Haenszel test",
    "tost": "Two One-Sided Tests for equivalence",
    "fleiss_kappa": "Fleiss' κ (multi-rater agreement)",
    "cohens_kappa": "Cohen's κ (inter-rater agreement)",
    "icc": "Intraclass correlation coefficient (Shrout & Fleiss 1979)",
    "bland_altman": "Bland-Altman analysis",
    "passing_bablok": "Passing-Bablok regression",
    "deming": "Deming regression",
    "concordance": "Lin's concordance correlation coefficient",
    "cronbach": "Cronbach's α reliability",
    "roc": "Receiver operating characteristic (ROC) analysis",
    "roc_compare": "DeLong's test for paired ROC AUCs",
    "calibration": "Logistic calibration (slope/intercept) + Hosmer-Lemeshow goodness-of-fit",
    "hosmer_lemeshow": "Hosmer-Lemeshow goodness-of-fit test",
    "dca": "Decision curve analysis (Vickers & Elkin 2006)",
    "linear": "Linear regression (OLS)",
    "logistic": "Multivariable logistic regression",
    "logistic_table": "Univariate + multivariable logistic regression (publication table)",
    "poisson": "Poisson regression",
    "gamma": "Gamma regression (GLM)",
    "negbinom": "Negative binomial regression",
    "rcs": "Restricted cubic spline regression",
    "cox_rcs": "Multivariable Cox-RCS regression",
    "ordinal": "Ordinal logistic regression (proportional odds)",
    "gee": "Generalized estimating equations (population-averaged)",
    "lmm": "Linear mixed model (random intercept / slope, REML)",
    "survival_km": "Kaplan-Meier survival analysis with log-rank test",
    "survival_cox": "Cox proportional hazards regression (+ Schoenfeld PH test)",
    "cox_tv": "Cox regression with time-varying covariates",
    "fine_gray": "Fine-Gray competing-risks regression",
    "landmark": "Landmark survival analysis",
    "evalue": "E-value sensitivity analysis (VanderWeele 2017)",
    "psm": "Propensity score matching (Austin 2011 standards)",
    "mice": "Multivariate imputation by chained equations (MICE)",
    "stepwise": "Stepwise variable selection (AIC / BIC / p)",
    "power": "Power / sample-size calculation",
    "table1": "Demographic baseline table (Table 1)",
}


class MethodAppendixRequest(BaseModel):
    session_id: str
    title: str = "Statistical Methods"
    include_data_io: bool = True       # include 'CSV / Excel imported, n rows, n cols'
    include_software: bool = True      # include software / package / version block


@router.post("/method_appendix")
def method_appendix_docx(req: MethodAppendixRequest):
    """Build a Methods-section DOCX from the session audit log.

    Walks the audit log, deduplicates analysis types, lists each in
    publication-ready phrasing alongside the underlying package + version
    + random seed (where applicable). Result is a one-shot 'Methods'
    section the user can paste into a manuscript.
    """
    if not HAS_DOCX:
        raise HTTPException(status_code=500, detail="python-docx not installed.")

    audit = store.get_audit(req.session_id)
    if not audit:
        raise HTTPException(status_code=404, detail="No audit log for this session (run an analysis first).")

    # Bucket by action code, count occurrences.
    counts: Dict[str, int] = {}
    seeds: List[int] = []
    for entry in audit:
        action = str(entry.get("action", ""))
        counts[action] = counts.get(action, 0) + 1
        params = entry.get("params") or {}
        seed = params.get("random_state") if isinstance(params, dict) else None
        if seed is not None and isinstance(seed, (int, float)):
            seeds.append(int(seed))

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    heading = doc.add_paragraph()
    run = heading.add_run(req.title)
    run.bold = True
    run.font.size = Pt(14)

    # ── Software paragraph ─────────────────────────────────────────────────
    if req.include_software:
        try:
            import sys as _sys
            import statsmodels as _sm
            import lifelines as _lf
            import sklearn as _sk
            import scipy as _sci
            import numpy as _np
            import pandas as _pd
            versions = (
                f"Python {_sys.version.split()[0]}, pandas {_pd.__version__}, NumPy {_np.__version__}, "
                f"SciPy {_sci.__version__}, statsmodels {_sm.__version__}, lifelines {_lf.__version__}, "
                f"scikit-learn {_sk.__version__}"
            )
        except Exception:
            versions = "Python 3.11+, pandas, NumPy, SciPy, statsmodels, lifelines, scikit-learn"
        seed_phrase = ""
        if seeds:
            uniq_seeds = sorted(set(seeds))
            seed_phrase = (
                f" Random-state seed{'s' if len(uniq_seeds) > 1 else ''} "
                f"used for reproducibility: {', '.join(str(s) for s in uniq_seeds)}."
            )
        doc.add_paragraph(
            f"All statistical analyses were performed in uSTAT (https://ustat.drtr.uk) using "
            f"{versions}.{seed_phrase} Two-sided α was set at 0.05 unless otherwise stated; "
            f"95% confidence intervals are reported for all effect estimates."
        )

    # ── Data I/O paragraph ─────────────────────────────────────────────────
    if req.include_data_io:
        df = store.get(req.session_id)
        if df is not None:
            n_rows, n_cols = df.shape
            doc.add_paragraph(
                f"The dataset comprised {n_rows} observations across {n_cols} variables. Missing values "
                f"were handled by listwise deletion unless an explicit imputation strategy (median or "
                f"multivariate imputation by chained equations) was selected for that analysis."
            )

    # ── Per-analysis methods bullets ───────────────────────────────────────
    para = doc.add_paragraph()
    para.add_run("Analyses performed").bold = True
    excluded = {"data_updated", "metadata_updated", "kind_override", "case_filter", "case_filter_cleared", "row_added", "row_deleted", "column_renamed", "computed_column"}
    used_actions = [a for a in counts.keys() if a and a not in excluded]
    seen_methods = set()
    for action in used_actions:
        human = _ACTION_HUMAN.get(action)
        if not human:
            continue  # skip unknown / non-analytic actions
        if human in seen_methods:
            continue
        seen_methods.add(human)
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(human)
        if counts[action] > 1:
            run = p.add_run(f"  (run {counts[action]} times)")
            run.italic = True

    if not seen_methods:
        doc.add_paragraph("No analyses recorded yet in the audit log.")

    # ── Citation footer ────────────────────────────────────────────────────
    foot = doc.add_paragraph()
    foot_run = foot.add_run(
        "\nCitation: Hoşoğlu Y. uSTAT — Browser-based Statistical Analysis Platform. "
        "https://ustat.drtr.uk (CITATION.cff in the repository for machine-readable form)."
    )
    foot_run.italic = True
    foot_run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        iter([buf.getvalue()]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="method_appendix.docx"'},
    )
