"""
AMA-style journal table formatter for uSTAT.
Converts raw Table 1 result data into publication-ready HTML, Excel, and Word.
"""
import io
import re
import numpy as np
import pandas as pd
from typing import Optional


# ═══════════════════════════════════════════════════════════════════════════════
# CONSTANTS
# ═══════════════════════════════════════════════════════════════════════════════

_TEST_SYMBOL_MAP = [
    (["student", "t-test", "t test", "independent t", "paired t"], "*"),
    (["mann", "whitney", "mann-whitney", "wilcoxon rank", "u test"], "\u2020"),
    (["chi-square", "chi square", "pearson chi"], "\u2021"),
    (["fisher"], "\u00A7"),
    (["anova", "one-way anova"], "\u00B6"),
    (["kruskal", "kruskal-wallis"], "**"),
    (["log-rank", "log rank", "mantel"], "\u2020\u2020"),
]

_TEST_DISPLAY_NAMES = {
    "*": "Student\u2019s t-test",
    "\u2020": "Mann-Whitney U test",
    "\u2021": "Chi-square test",
    "\u00A7": "Fisher\u2019s exact test",
    "\u00B6": "One-way ANOVA",
    "**": "Kruskal-Wallis test",
    "\u2020\u2020": "Log-rank test",
}

_KNOWN_ABBREVIATIONS = {
    "NYHA": "New York Heart Association", "LVEF": "left ventricular ejection fraction",
    "IQR": "interquartile range", "SD": "standard deviation", "BMI": "body mass index",
    "MACE": "major adverse cardiovascular events", "MI": "myocardial infarction",
    "CV": "cardiovascular", "HF": "heart failure", "CKD": "chronic kidney disease",
    "eGFR": "estimated glomerular filtration rate", "BNP": "B-type natriuretic peptide",
    "NT-proBNP": "N-terminal pro-B-type natriuretic peptide",
    "ACEI": "angiotensin-converting enzyme inhibitor", "ARB": "angiotensin receptor blocker",
    "ARNI": "angiotensin receptor-neprilysin inhibitor", "MRA": "mineralocorticoid receptor antagonist",
    "SGLT2i": "sodium-glucose cotransporter 2 inhibitor",
    "CABG": "coronary artery bypass grafting", "PCI": "percutaneous coronary intervention",
    "AF": "atrial fibrillation", "COPD": "chronic obstructive pulmonary disease",
    "DM": "diabetes mellitus", "HTN": "hypertension",
    "LDL": "low-density lipoprotein", "HDL": "high-density lipoprotein",
    "TG": "triglycerides", "HbA1c": "glycated hemoglobin",
    "CI": "confidence interval", "HR": "hazard ratio", "OR": "odds ratio",
    "AE": "adverse event", "SAE": "serious adverse event", "ITT": "intention-to-treat",
    "KM": "Kaplan-Meier", "STEMI": "ST-elevation myocardial infarction",
    "NSTEMI": "non-ST-elevation myocardial infarction", "ACS": "acute coronary syndrome",
    "CAD": "coronary artery disease", "PAD": "peripheral artery disease",
    "TIMI": "Thrombolysis in Myocardial Infarction",
    "BARC": "Bleeding Academic Research Consortium",
    "PNI": "prognostic nutritional index", "SII": "systemic immune-inflammation index",
    "TyG": "triglyceride-glucose index", "NLR": "neutrophil-to-lymphocyte ratio",
    "PLR": "platelet-to-lymphocyte ratio", "CRP": "C-reactive protein",
    "GFR": "glomerular filtration rate", "EF": "ejection fraction",
}

_UNIT_PATTERNS = {
    r'\bage\b|\bya[sş]\b': "years",
    r'\bheight\b|\bboy\b': "cm", r'\bweight\b|\bk[ıi]lo\b': "kg",
    r'\bbmi\b|\bvk[iİ]\b': "kg/m\u00B2",
    r'\bblood pressure\b|\bsbp\b|\bdbp\b|\bsystolic\b|\bdiastolic\b': "mmHg",
    r'\bheart rate\b|\bnab[ıi]z\b': "bpm",
    r'\bcreatinine?\b|\bkreatinin\b': "mg/dL",
    r'\begfr\b|\bgfr\b': "mL/min/1.73m\u00B2",
    r'\bhemoglobin\b|\bhgb\b|\bhb\b': "g/dL",
    r'\bplatelet\b|\bplt\b': "\u00D710\u00B3/\u03BCL",
    r'\bwbc\b|\bleukocyte?\b': "\u00D710\u00B3/\u03BCL",
    r'\bcholesterol\b|\bldl\b|\bhdl\b': "mg/dL",
    r'\btriglyceride?\b|\btg\b': "mg/dL",
    r'\bglucose?\b|\bglukoz\b': "mg/dL", r'\bhba1c\b': "%",
    r'\blvef\b|\bejection fraction\b': "%",
    r'\bnt-probnp\b|\bbnp\b': "pg/mL", r'\btroponin\b': "ng/mL",
    r'\bcrp\b|\bc-reactive\b': "mg/L",
    r'\balbumin\b': "g/dL", r'\bsodium\b|\bna\b': "mEq/L", r'\bpotassium\b|\bk\b': "mEq/L",
    r'\bast\b|\bsgot\b': "U/L", r'\balt\b|\bsgpt\b': "U/L",
}


# ═══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def _detect_unit(var_name: str) -> Optional[str]:
    name_lower = var_name.lower()
    for pattern, unit in _UNIT_PATTERNS.items():
        if re.search(pattern, name_lower):
            # Skip if name already contains the unit
            if unit.lower().replace("\u00B2", "2") in name_lower.replace("\u00B2", "2"):
                return None
            return unit
    return None


def _assign_test_symbol(test_name: str) -> str:
    if not test_name:
        return ""
    t = test_name.lower()
    for keywords, symbol in _TEST_SYMBOL_MAP:
        if any(kw in t for kw in keywords):
            return symbol
    return ""


def _find_abbreviations(text: str) -> dict:
    found = set(re.findall(r'\b([A-Z][A-Za-z]{1,8})\b', text))
    found.update(re.findall(r'\b([a-z]?[A-Z][A-Za-z]{1,8})\b', text))
    return {a: _KNOWN_ABBREVIATIONS[a] for a in sorted(found) if a in _KNOWN_ABBREVIATIONS}


def _fmt_p(p) -> str:
    if p is None or p == "" or p == "N/A":
        return str(p) if p else ""
    s = str(p).strip()
    if s.startswith("<"):
        return s  # already formatted like "<0.001"
    try:
        pv = float(s)
    except (ValueError, TypeError):
        return s
    if pv < 0.001:
        return "<0.001"
    if pv < 0.01:
        return f"{pv:.3f}"
    return f"{pv:.2f}"


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN FORMAT FUNCTION
# ═══════════════════════════════════════════════════════════════════════════════

def format_table1_for_journal(result: dict, options: dict = None) -> dict:
    """Convert uSTAT Table 1 result into AMA journal-formatted output.

    Args:
        result: The dict returned by /api/stats/table1
        options: {bold_significant_p: bool, show_test_column: bool, table_number: int}

    Returns:
        {title, columns, rows, html, footnotes, abbreviations, validation}
    """
    opts = options or {}
    bold_p = opts.get("bold_significant_p", True)
    show_test = opts.get("show_test_column", False)
    table_num = opts.get("table_number", 1)

    group_col = result.get("group_column", "")
    group_labels = result.get("group_labels", [])
    group_ns = result.get("group_ns", {})
    total_n = result.get("total_n", 0)
    has_groups = len(group_labels) >= 2

    # Build column headers
    columns = ["Variable"]
    if has_groups:
        for g in group_labels:
            n = group_ns.get(g, "")
            columns.append(f"{g} (n={n})")
    else:
        columns.append(f"Overall (n={total_n})")
    if has_groups:
        columns.append("p-value")

    # Build formatted rows
    formatted_rows = []
    used_tests = set()
    all_text = ""

    for row in result.get("rows", []):
        var_name = row.get("variable", "")
        all_text += " " + var_name

        if row["type"] == "numeric":
            # stat_rows: [{label, overall, group_stats: {group_label: value}}]
            stat_rows = row.get("stat_rows", [])
            for si, sr in enumerate(stat_rows):
                stat_label = sr.get("label", "")

                if si == 0:
                    # Build label: "Variable, unit, stat_format"
                    label = var_name
                    unit = _detect_unit(var_name)
                    if unit:
                        label = f"{var_name}, {unit}"
                    if stat_label:
                        label += f", {stat_label}"
                else:
                    label = f"\u2003{stat_label}" if stat_label else ""

                # Extract values from group_stats dict
                values = []
                if has_groups:
                    grp = sr.get("group_stats", {})
                    for g in group_labels:
                        values.append(str(grp.get(g, "")))
                else:
                    values.append(str(sr.get("overall", "")))

                p_val = ""
                test_sym = ""
                if si == 0 and has_groups:
                    p_val = _fmt_p(row.get("p_value"))
                    test_name = row.get("test", "") or ""
                    test_sym = _assign_test_symbol(test_name)
                    if test_sym:
                        used_tests.add(test_sym)

                formatted_rows.append({
                    "label": label,
                    "values": values,
                    "p_value": p_val,
                    "test_symbol": test_sym,
                    "indent": 1 if si > 0 else 0,
                    "bold": False,
                })

        elif row["type"] == "categorical":
            # p-value and test for header row
            p_val_cat = _fmt_p(row.get("p_value")) if has_groups else ""
            test_name_cat = row.get("test", "") or ""
            test_sym_cat = _assign_test_symbol(test_name_cat) if has_groups else ""
            if test_sym_cat:
                used_tests.add(test_sym_cat)

            # Header row (variable name)
            formatted_rows.append({
                "label": var_name + ", n (%)",
                "values": [""] * (len(group_labels) if has_groups else 1),
                "p_value": p_val_cat,
                "test_symbol": test_sym_cat,
                "indent": 0,
                "bold": True,
            })

            # Category sub-rows from sub_rows: [{category, overall, group_stats: {gl: "n (pct%)"}}]
            for cat in row.get("sub_rows", []):
                cat_label = str(cat.get("category", ""))
                values = []
                if has_groups:
                    grp = cat.get("group_stats", {})
                    for g in group_labels:
                        values.append(str(grp.get(g, "")))
                else:
                    values.append(str(cat.get("overall", "")))

                formatted_rows.append({
                    "label": cat_label,
                    "values": values,
                    "p_value": "",
                    "test_symbol": "",
                    "indent": 1,
                    "bold": False,
                })

    # Detect abbreviations
    abbreviations = _find_abbreviations(all_text)

    # Build footnotes
    footnotes = []
    if abbreviations:
        abbr_str = "; ".join(f"{k} = {v}" for k, v in abbreviations.items())
        footnotes.append(f"Abbreviations: {abbr_str}")
    if used_tests:
        test_notes = "; ".join(f"{sym} {_TEST_DISPLAY_NAMES.get(sym, '')}" for sym in sorted(used_tests))
        footnotes.append(f"Statistical tests: {test_notes}")
    footnotes.append("Values are presented as mean \u00B1 SD, median [IQR], or n (%) as appropriate.")

    # Title
    title = f"Table {table_num}. Baseline Clinical Characteristics"
    if group_col:
        title += f" by {group_col}"

    # HTML
    html = _generate_html(title, columns, formatted_rows, footnotes, has_groups, bold_p)

    # Validation
    validation = _validate_table(formatted_rows, columns, footnotes, has_groups, abbreviations)

    return {
        "title": title,
        "columns": columns,
        "rows": formatted_rows,
        "html": html,
        "footnotes": footnotes,
        "abbreviations": abbreviations,
        "validation": validation,
    }


# ═══════════════════════════════════════════════════════════════════════════════
# HTML GENERATION
# ═══════════════════════════════════════════════════════════════════════════════

def _generate_html(title, columns, rows, footnotes, has_p, bold_p=True):
    css = """<style>
.journal-table { border-collapse: collapse; width: 100%; font-family: 'Times New Roman', Georgia, serif; font-size: 10pt; line-height: 1.4; }
.journal-table caption { text-align: left; font-weight: bold; font-size: 11pt; padding-bottom: 6pt; caption-side: top; }
.journal-table thead tr { border-top: 2pt solid #000; border-bottom: 0.75pt solid #000; }
.journal-table thead th { text-align: center; font-weight: bold; padding: 6pt 8pt; vertical-align: bottom; background: white; border: none; }
.journal-table thead th:first-child { text-align: left; }
.journal-table tbody td { padding: 3pt 8pt; border: none; background: white; text-align: center; vertical-align: top; }
.journal-table tbody td:first-child { text-align: left; }
.journal-table tbody tr:last-child { border-bottom: 2pt solid #000; }
.journal-table .indent-1 { padding-left: 2em; }
.journal-table .bold-row { font-weight: bold; }
.journal-table .p-sig { font-weight: bold; }
.journal-table sup.test-sym { font-size: 7pt; vertical-align: super; color: #444; margin-left: 1pt; }
.journal-table tfoot td { border: none; padding-top: 4pt; font-size: 9pt; color: #333; text-align: left; line-height: 1.5; font-style: italic; }
</style>"""

    html = f'{css}\n<table class="journal-table">\n  <caption>{title}</caption>\n  <thead>\n    <tr>\n'
    for col in columns:
        html += f'      <th>{col}</th>\n'
    html += '    </tr>\n  </thead>\n  <tbody>\n'

    for row in rows:
        cls_parts = []
        if row["indent"] > 0:
            cls_parts.append("indent-1")
        if row["bold"]:
            cls_parts.append("bold-row")
        cls = f' class="{" ".join(cls_parts)}"' if cls_parts else ""

        label = f"<strong>{row['label']}</strong>" if row["bold"] else row["label"]
        html += f'    <tr>\n      <td{cls}>{label}</td>\n'

        for val in row["values"]:
            html += f'      <td>{val}</td>\n'

        if has_p:
            p_val = row.get("p_value", "")
            sym = row.get("test_symbol", "")
            is_sig = False
            if p_val == "<0.001":
                is_sig = True
            elif p_val:
                try:
                    is_sig = float(p_val) < 0.05
                except ValueError:
                    pass
            p_display = f'{p_val}<sup class="test-sym">{sym}</sup>' if sym and p_val else p_val
            p_cls = ' class="p-sig"' if (is_sig and bold_p) else ""
            html += f'      <td{p_cls}>{p_display}</td>\n'

        html += '    </tr>\n'

    html += '  </tbody>\n'
    if footnotes:
        notes = "<br>".join(footnotes)
        html += f'  <tfoot>\n    <tr>\n      <td colspan="{len(columns)}">{notes}</td>\n    </tr>\n  </tfoot>\n'
    html += '</table>'
    return html


# ═══════════════════════════════════════════════════════════════════════════════
# VALIDATION
# ═══════════════════════════════════════════════════════════════════════════════

def _validate_table(rows, columns, footnotes, has_p, abbreviations):
    checks = {}
    checks["row_formatting"] = "PASS" if any(r["indent"] > 0 for r in rows) else "WARN"

    p_ok = True
    for r in rows:
        p = r.get("p_value", "")
        if not p or p == "<0.001":
            continue
        try:
            pv = float(p)
            decimals = len(p.split(".")[-1]) if "." in p else 0
            if pv >= 0.05 and decimals > 2:
                p_ok = False
            if 0.001 <= pv < 0.05 and decimals < 3:
                p_ok = False
        except ValueError:
            pass
    checks["p_value_formatting"] = "PASS" if p_ok else "FAIL"
    checks["abbreviation_footnote"] = "PASS" if (not abbreviations or any("Abbreviation" in f for f in footnotes)) else "FAIL"
    checks["statistical_test_footnote"] = "PASS" if (not has_p or any("Statistical" in f for f in footnotes)) else "FAIL"
    checks["table_title"] = "PASS"
    checks["no_empty_rows"] = "PASS" if not any(not r["label"].strip() for r in rows) else "FAIL"

    checks["status"] = "READY FOR SUBMISSION" if all(v == "PASS" for v in checks.values()) else "REVISION REQUIRED"
    return checks


# ═══════════════════════════════════════════════════════════════════════════════
# EXPORT: EXCEL
# ═══════════════════════════════════════════════════════════════════════════════

def export_journal_excel(formatted: dict) -> bytes:
    from openpyxl.styles import Font, Alignment, Border, Side

    buf = io.BytesIO()
    title = formatted["title"]
    columns = formatted["columns"]
    rows = formatted["rows"]
    footnotes = formatted.get("footnotes", [])

    data = []
    for row in rows:
        label = ("\u2003" + row["label"]) if row.get("indent", 0) > 0 else row["label"]
        r = [label] + row.get("values", [])
        if "p-value" in columns:
            p = row.get("p_value", "")
            sym = row.get("test_symbol", "")
            r.append(f"{p}{sym}" if sym and p else p)
        data.append(r)

    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        df_out = pd.DataFrame(data, columns=columns)
        df_out.to_excel(writer, sheet_name="Table", index=False, startrow=1)
        ws = writer.sheets["Table"]

        ws.cell(row=1, column=1, value=title)
        ws.cell(row=1, column=1).font = Font(bold=True, size=11)

        thin = Side(style='thin')
        thick = Side(style='medium')
        for col_idx in range(1, len(columns) + 1):
            cell = ws.cell(row=2, column=col_idx)
            cell.font = Font(bold=True, size=10)
            cell.alignment = Alignment(horizontal='center' if col_idx > 1 else 'left')
            cell.border = Border(top=thick, bottom=thin)

        for row_idx, row in enumerate(rows, start=3):
            for col_idx in range(1, len(columns) + 1):
                cell = ws.cell(row=row_idx, column=col_idx)
                cell.alignment = Alignment(horizontal='center' if col_idx > 1 else 'left')
                if row.get("bold"):
                    cell.font = Font(bold=True, size=10)
                if row_idx == len(rows) + 2:
                    cell.border = Border(bottom=thick)

        start = len(data) + 3
        for i, fn in enumerate(footnotes):
            ws.cell(row=start + i, column=1, value=fn)
            ws.cell(row=start + i, column=1).font = Font(italic=True, size=9)

    buf.seek(0)
    return buf.read()


# ═══════════════════════════════════════════════════════════════════════════════
# EXPORT: WORD (DOCX)
# ═══════════════════════════════════════════════════════════════════════════════

def export_journal_word(formatted: dict) -> bytes:
    from docx import Document as DocxDocument
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from lxml import etree

    doc = DocxDocument()
    title = formatted["title"]
    columns = formatted["columns"]
    rows = formatted["rows"]
    footnotes = formatted.get("footnotes", [])

    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)

    n_cols = len(columns)
    n_rows = len(rows) + 1
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    THICK = {'val': 'single', 'sz': 12, 'color': '000000'}
    THIN = {'val': 'single', 'sz': 4, 'color': '000000'}

    def _set_border(cell, **kwargs):
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        existing = tcPr.find(qn('w:tcBorders'))
        if existing is not None:
            tcPr.remove(existing)
        tcBorders = etree.SubElement(tcPr, qn('w:tcBorders'))
        for side in ['top', 'left', 'bottom', 'right']:
            spec = kwargs.get(side)
            el = etree.SubElement(tcBorders, qn(f'w:{side}'))
            if spec:
                for k, v in spec.items():
                    el.set(qn(f'w:{k}'), str(v))
            else:
                el.set(qn('w:val'), 'none')

    # Header
    for j, col_name in enumerate(columns):
        cell = table.rows[0].cells[j]
        cell.text = col_name
        _set_border(cell, top=THICK, bottom=THIN)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data rows
    for i, row_data in enumerate(rows):
        label = ("\u2003" + row_data["label"]) if row_data.get("indent", 0) > 0 else row_data["label"]
        values = [label] + row_data.get("values", [])
        if "p-value" in columns:
            p_val = row_data.get("p_value", "")
            sym = row_data.get("test_symbol", "")
            values.append(f"{p_val}{sym}" if sym and p_val else p_val)

        for j in range(min(n_cols, len(values))):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(values[j]) if values[j] is not None else ""
            border_spec = {}
            if i == len(rows) - 1:
                border_spec['bottom'] = THICK
            _set_border(cell, **border_spec)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.size = Pt(10)
                    if row_data.get("bold"):
                        run.bold = True

    if footnotes:
        p = doc.add_paragraph()
        for fn in footnotes:
            run = p.add_run(fn + "\n")
            run.italic = True
            run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
